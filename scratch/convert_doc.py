import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def main():
    md_path = "CLIENT_CONTENT_TEMPLATE.md"
    docx_path = "CLIENT_CONTENT_TEMPLATE.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found!")
        return
        
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styling variables
    COLOR_PRIMARY = RGBColor(42, 27, 21)   # Dark Cocoa/Charcoal
    COLOR_ACCENT = RGBColor(220, 38, 38)   # Primary red
    COLOR_MUTED = RGBColor(139, 126, 116)  # Warm Muted
    FONT_NAME = 'Arial'

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("ORBIKO STUDIO")
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("CLIENT CONTENT COLLECTION DOCUMENT")
    subtitle_run.font.name = FONT_NAME
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = COLOR_ACCENT
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Read Markdown and parse simply
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    
    for line in lines:
        stripped = line.strip()
        
        # Code block tracking
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line.rstrip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_PRIMARY
            continue
            
        # Headings
        if stripped.startswith("# "):
            title = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            run.font.name = FONT_NAME
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = COLOR_PRIMARY
            
        elif stripped.startswith("## "):
            title = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            run.font.name = FONT_NAME
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = COLOR_PRIMARY
            
        elif stripped.startswith("### "):
            title = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = COLOR_ACCENT
            
        # Lists
        elif stripped.startswith("* "):
            item = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(item)
            run.font.name = FONT_NAME
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_PRIMARY
            
        elif stripped.startswith("- "):
            item = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(item)
            run.font.name = FONT_NAME
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_PRIMARY
            
        elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. ") or stripped.startswith("4. ") or stripped.startswith("5. ") or stripped.startswith("6. "):
            item = stripped[3:]
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(item)
            run.font.name = FONT_NAME
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_PRIMARY
            
        # Dividers
        elif stripped == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("____________________________________________________")
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_MUTED
            
        # Regular text / forms lines
        elif stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            # Highlight labels ending in ":"
            if ":" in stripped:
                parts = stripped.split(":", 1)
                run_label = p.add_run(parts[0] + ":")
                run_label.font.name = FONT_NAME
                run_label.font.size = Pt(10.5)
                run_label.font.bold = True
                run_label.font.color.rgb = COLOR_PRIMARY
                
                if parts[1].strip():
                    run_text = p.add_run(parts[1])
                    run_text.font.name = FONT_NAME
                    run_text.font.size = Pt(10.5)
                    run_text.font.color.rgb = COLOR_PRIMARY
                else:
                    # Provide an empty underline space for client response
                    run_blank = p.add_run(" __________________________________________________")
                    run_blank.font.name = FONT_NAME
                    run_blank.font.size = Pt(10.5)
                    run_blank.font.color.rgb = COLOR_MUTED
            else:
                run = p.add_run(stripped)
                run.font.name = FONT_NAME
                run.font.size = Pt(10.5)
                run.font.color.rgb = COLOR_PRIMARY
        else:
            # Empty line spacer
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
    doc.save(docx_path)
    print("Successfully generated Word Doc!")

if __name__ == "__main__":
    main()
